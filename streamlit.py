import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi
import google.generativeai as genai
from io import BytesIO
import os
from dotenv import load_dotenv
from docx import Document
from markdown2 import markdown
from bs4 import BeautifulSoup
from urllib.parse import urlparse, parse_qs

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

if not api_key:
    st.error("Google API key not found. Please set the GOOGLE_API_KEY environment variable.")
    st.stop()

genai.configure(api_key=api_key)

# Define prompt for summarization
prompt = """You are a YouTube video summarizer. You will be taking the transcript text
and summarizing the entire video and providing the important summary in points
within 2500 words. Please provide the summary of the text given here:  """

# Function to extract transcript details
def extract_transcript_details(youtube_video_url):
    try:
        video_id = get_video_id(youtube_video_url)
        if not video_id:
            raise ValueError("Invalid YouTube URL")
        transcript_text = YouTubeTranscriptApi.get_transcript(video_id)
        transcript = " ".join(item["text"] for item in transcript_text)
        return transcript
    except Exception as e:
        st.error(f"An error occurred while extracting the transcript: {e}")
        return None

# Function to generate content summary using Google Gemini Pro
def generate_gemini_content(transcript_text, prompt):
    try:
        model = genai.GenerativeModel("gemini-1.5-pro-latest")
        response = model.generate_content(prompt + transcript_text)
        return response.text
    except Exception as e:
        st.error(f"An error occurred while generating the summary: {e}")
        return None

# Function to generate DOCX from Markdown summary text
def generate_docx(summary):
    doc = Document()
    doc.add_heading('Detailed Notes', level=1)
    html_summary = markdown(summary)
    soup = BeautifulSoup(html_summary, 'html.parser')

    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'strong':
            doc.add_paragraph(element.text, style='Bold')
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')

    docx_output = BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)
    
    return docx_output

# Helper function to extract video ID from URL
def get_video_id(url):
    parsed_url = urlparse(url)
    video_id = parse_qs(parsed_url.query).get('v')
    return video_id[0] if video_id else None

# Streamlit app
st.markdown("""
    <h1 style='font-size: 32px;'>Transform YouTube 📹 Transcripts into Impactful Detailed Notes</h1>
    """, unsafe_allow_html=True)

# Initialize or clear session state
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'docx_file' not in st.session_state:
    st.session_state.docx_file = None

def reset_app():
    st.session_state.summary = None
    st.session_state.docx_file = None
    st.rerun()

youtube_link = st.text_input("Enter YouTube Video Link:")

if youtube_link:
    video_id = get_video_id(youtube_link)
    if video_id:
        st.image(f"http://img.youtube.com/vi/{video_id}/0.jpg", use_column_width=True)
    
    if st.button("Get Detailed Notes"):
        with st.spinner("Extracting transcript and generating summary..."):
            transcript_text = extract_transcript_details(youtube_link)

            if transcript_text:
                st.session_state.summary = generate_gemini_content(transcript_text, prompt)
                
                if st.session_state.summary:
                    st.markdown("## Detailed Notes:")
                    st.write(st.session_state.summary)

                    # Generate DOCX
                    st.session_state.docx_file = generate_docx(st.session_state.summary)
                    
                else:
                    st.error("Failed to generate summary.")
            else:
                st.error("Failed to extract transcript.")

# Add the download button to the sidebar
st.sidebar.markdown("Enter the YouTube Video Link and Click Get Detailed Notes button.")
st.sidebar.title("DOWNLOAD MENU")

with st.sidebar:
    if st.session_state.docx_file:
        st.markdown("Summary Generated Successfully !!!")
        st.write(" ")
        st.info("Click the button below to download your detailed notes in DOCX format.")
        st.download_button(
            label="Download Summary and Exit",
            data=st.session_state.docx_file,
            file_name="summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.write("No summary available for download.")


st.sidebar.markdown("""
    <h1 style='font-size: 20px;'>Manage Reset</h1>""", unsafe_allow_html=True)
if st.sidebar.button("Reset"):
    
    reset_app()

 # Add the credit section
st.sidebar.markdown("<hr>", unsafe_allow_html=True)  # Adds a horizontal line with HTML
st.sidebar.markdown("<h3 style='color: #2ca02c;font-size: 20px;'>App Created by: Parthebhan Pari</h3>", unsafe_allow_html=True)